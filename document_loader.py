import os
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pinecone import Pinecone, ServerlessSpec

class DocumentManager:
    def __init__(self):
        """Pinecone bağlantısını kurar."""
        api_key = os.environ.get("PINECONE_API_KEY")
        self.pc = Pinecone(api_key=api_key)

        index_name = "mythority-memory"

        existing_indexes = [i.name for i in self.pc.list_indexes()]
        if index_name not in existing_indexes:
            print(f"Sistem: '{index_name}' index'i oluşturuluyor...")
            self.pc.create_index(
                name=index_name,
                dimension=1024,  # multilingual-e5-large boyutu
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )

        self.index = self.pc.Index(index_name)
        # Pinecone'un kendi inference API'si — model RAM'e yüklenmiyor
        self.embed_model = "multilingual-e5-large"
        print("Sistem: Pinecone vektör veritabanı aktif.\n")

    def get_embeddings(self, texts):
        """Pinecone inference API ile embedding üretir."""
        result = self.pc.inference.embed(
            model=self.embed_model,
            inputs=texts,
            parameters={"input_type": "passage"}
        )
        return [item["values"] for item in result]

    def read_file(self, file_path):
        """Dosya uzantısına göre doğru okuyucuyu seçer ve metni döndürür."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            from pypdf import PdfReader
            text = ""
            reader = PdfReader(file_path)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text

        elif ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text.strip()
            return text

        elif ext == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            text = ""
            for sheet in wb.worksheets:
                text += f"[Sayfa: {sheet.title}]\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(c) for c in row if c is not None])
                    if row_text.strip():
                        text += row_text + "\n"
            return text

        elif ext in {".txt", ".py", ".js", ".ts", ".html", ".css", ".java", ".cpp",
                     ".c", ".cs", ".php", ".rb", ".go", ".rs", ".swift", ".kt",
                     ".json", ".xml", ".yaml", ".yml", ".md", ".sh", ".bat"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return f"[Dosya: {os.path.basename(file_path)}]\n\n{content}"

        else:
            raise ValueError(f"Desteklenmeyen dosya formatı: {ext}")

    def sanitize_id(self, text):
        """Türkçe ve özel karakterleri ASCII'ye çevirir."""
        import unicodedata
        normalized = unicodedata.normalize("NFKD", text)
        ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
        # Boşluk ve özel karakterleri alt çizgiye çevir
        return "".join(c if c.isalnum() or c in "-_." else "_" for c in ascii_text)

    def add_document(self, file_path, user_id):
        """Dosyayı okur, parçalar, vektöre çevirir ve Pinecone'a ekler."""
        text = self.read_file(file_path)
        if not text.strip():
            raise ValueError("Dosyadan metin okunamadı veya dosya boş.")

        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = splitter.split_text(text)

        embeddings = self.get_embeddings(chunks)

        vectors = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            vectors.append({
                "id": f"{self.sanitize_id(user_id)}_{self.sanitize_id(os.path.basename(file_path))}_chunk_{i}",
                "values": embedding,
                "metadata": {
                    "owner": user_id,
                    "text": chunk,
                    "filename": os.path.basename(file_path)
                }
            })

        batch_size = 100
        for i in range(0, len(vectors), batch_size):
            self.index.upsert(vectors=vectors[i:i + batch_size])

        print(f"Sistem: {os.path.basename(file_path)} dosyası {user_id} hafızasına kaydedildi.")

    def search_memory(self, query, user_id, n_results=3):
        """Kullanıcıya ait belgeler arasında arama yapar."""
        try:
            result = self.pc.inference.embed(
                model=self.embed_model,
                inputs=[query],
                parameters={"input_type": "query"}
            )
            query_embedding = result[0]["values"]

            results = self.index.query(
                vector=query_embedding,
                top_k=n_results,
                filter={"owner": {"$eq": user_id}},
                include_metadata=True
            )
            return [match["metadata"]["text"] for match in results["matches"]]
        except Exception as e:
            print(f"Arama hatası: {e}")
            return []

    def get_all_documents(self, user_id):
        """Kullanıcıya ait benzersiz dosya isimlerini getirir."""
        try:
            dummy = self.pc.inference.embed(
                model=self.embed_model,
                inputs=["_"],
                parameters={"input_type": "query"}
            )
            results = self.index.query(
                vector=dummy[0]["values"],
                top_k=1000,
                filter={"owner": {"$eq": user_id}},
                include_metadata=True
            )
            return list({m["metadata"].get("filename", "") for m in results["matches"] if m["metadata"].get("filename")})
        except Exception as e:
            print(f"Hafıza okuma hatası: {e}")
            return []

    def clear_memory(self):
        """Tüm vektörel hafızayı sıfırlar."""
        try:
            self.index.delete(delete_all=True)
            print("Sistem: Pinecone hafızası tamamen sıfırlandı.")
            return True
        except Exception as e:
            print(f"Sıfırlama hatası: {e}")
            return False